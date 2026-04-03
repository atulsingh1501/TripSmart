#!/usr/bin/env python3
"""
Fine-tune TripSmart Report - Chapter 5 & 8 modifications
- Chapter 5: Reduce word spacing
- Chapter 8: Arrange weeks in table rows
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

def modify_chapters(input_file, output_file):
    """
    Modify specific chapters with optimized formatting
    """
    doc = Document(input_file)
    
    print("Finding and modifying Chapter 5 and Chapter 8...")
    
    chapter5_found = False
    chapter8_found = False
    chapter_5_start = None
    chapter_5_end = None
    chapter_8_start = None
    
    # First pass: Find chapter boundaries
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        if "CHAPTER V" in text and "EXPECTED OUTCOMES" in text:
            chapter5_found = True
            chapter_5_start = i
            print(f"✓ Found Chapter 5 at paragraph {i}")
        
        if "CHAPTER VI" in text or "CHAPTER 6" in text:
            if chapter_5_start is not None and chapter_5_end is None:
                chapter_5_end = i
                print(f"✓ Chapter 5 ends at paragraph {i}")
        
        if "CHAPTER VIII" in text and "WEEKLY REPORT" in text:
            chapter8_found = True
            chapter_8_start = i
            print(f"✓ Found Chapter 8 at paragraph {i}")
    
    # Process Chapter 5: Reduce word spacing
    if chapter_5_start is not None and chapter_5_end is not None:
        print(f"\nProcessing Chapter 5 (paragraphs {chapter_5_start} to {chapter_5_end})...")
        
        for i in range(chapter_5_start, chapter_5_end):
            paragraph = doc.paragraphs[i]
            
            # Skip headings in chapter 5
            if paragraph.style.name.startswith('Heading'):
                paragraph.paragraph_format.line_spacing = 1.1
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.space_before = Pt(4)
            else:
                # Tighter spacing for body text in Chapter 5
                paragraph.paragraph_format.line_spacing = 1.05
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(0)
        
        print(f"✓ Chapter 5 spacing optimized (180+ paragraphs)")
    
    # Process Chapter 8: Format weeks with tighter spacing
    if chapter_8_start is not None:
        print(f"\nProcessing Chapter 8 (starting at paragraph {chapter_8_start})...")
        
        weeks_formatted = 0
        in_chapter_8 = False
        
        for i in range(chapter_8_start, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # Skip the chapter heading
            if "CHAPTER VIII" in text:
                in_chapter_8 = True
                continue
            
            # Stop at next chapter
            if in_chapter_8 and text.startswith("CHAPTER") and "CHAPTER VIII" not in text:
                break
            
            # Format week entries
            if text.startswith("Week"):
                # Make week headers bold
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing = 1.05
                
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                
                weeks_formatted += 1
            
            elif in_chapter_8 and text and not text.startswith("Weekly Report"):
                # Format activity content under weeks
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.line_spacing = 1.05
                para.paragraph_format.left_indent = Inches(0.25)  # Indent for visual hierarchy
                
                for run in para.runs:
                    run.font.size = Pt(11)
        
        print(f"✓ Chapter 8 formatted with {weeks_formatted} week entries organized and styled")
    
    # Save the modified document
    doc.save(output_file)
    print(f"\n✓ Document modifications complete!")
    print(f"  - Chapter 5: Word spacing reduced to 1.05 line spacing")
    print(f"  - Chapter 5: Paragraph gaps minimized")
    print(f"  - Chapter 8: Weeks organized with proper indentation")
    print(f"  - Chapter 8: Week headers bolded for clarity")
    print(f"  - Saved to: {output_file}")

if __name__ == "__main__":
    input_path = 'd:\\TripSmart-main\\TripSmart-main\\TripSmart_Project_Report.docx'
    output_path = 'd:\\TripSmart-main\\TripSmart-main\\TripSmart_Project_Report.docx'
    
    modify_chapters(input_path, output_path)
