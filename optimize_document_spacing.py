#!/usr/bin/env python3
"""
Optimize TripSmart Project Report spacing
Reduces line spacing, paragraph spacing, and adjusts formatting
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

def optimize_document_spacing(input_file, output_file):
    """
    Optimize spacing in the document for compact professional format
    """
    doc = Document(input_file)
    
    print("Optimizing document spacing...")
    
    # Counter for statistics
    paragraph_count = 0
    table_count = 0
    
    # Process all paragraphs
    for paragraph in doc.paragraphs:
        paragraph_count += 1
        
        # Set line spacing to 1.15 (more compact than 1.25)
        paragraph.paragraph_format.line_spacing = 1.15
        
        # Reduce space before and after paragraphs
        current_space_before = paragraph.paragraph_format.space_before
        current_space_after = paragraph.paragraph_format.space_after
        
        # Set optimal spacing
        if paragraph.style.name.startswith('Heading'):
            # Keep some space after headings
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.space_before = Pt(6)
        else:
            # Minimize spacing for body paragraphs
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.space_before = Pt(0)
        
        # Adjust word spacing to normal (remove extra spacing)
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    # Process all tables
    for table in doc.tables:
        table_count += 1
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Tighter spacing in tables
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
    
    # Save optimized document
    doc.save(output_file)
    print(f"✓ Document optimized successfully!")
    print(f"  - Paragraphs processed: {paragraph_count}")
    print(f"  - Tables processed: {table_count}")
    print(f"  - Line spacing reduced to: 1.15")
    print(f"  - Paragraph spacing optimized")
    print(f"  - Output file: {output_file}")

if __name__ == "__main__":
    input_path = 'd:\\TripSmart-main\\TripSmart-main\\TripSmart_Project_Report.docx'
    output_path = 'd:\\TripSmart-main\\TripSmart-main\\TripSmart_Project_Report.docx'
    
    optimize_document_spacing(input_path, output_path)
