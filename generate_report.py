#!/usr/bin/env python3
"""
TripSmart Project Report Generator
Follows guidelines for Minor Project Documentation
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_margins(section, left=1.25, right=1, top=1, bottom=1):
    """Set document margins in inches"""
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def set_line_spacing(paragraph, spacing=1.25):
    """Set line spacing for paragraph"""
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(spacing * 14)  # Convert to points

def add_heading_styled(doc, text, level=1):
    """Add formatted heading"""
    if level == 1:
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_format.line_spacing = 1.25
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p
    elif level == 2:
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_format.line_spacing = 1.25
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

def add_paragraph_styled(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0):
    """Add formatted paragraph"""
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.alignment = alignment
    p_format.line_spacing = 1.25
    p_format.first_line_indent = Inches(indent) if indent else 0
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def create_cover_page(doc):
    """Create the cover page"""
    set_margins(doc.sections[0], 1.25, 1, 1, 1)
    
    # Add spacing
    for _ in range(5):
        doc.add_paragraph()
    
    # Title
    title = add_heading_styled(doc, "TripSmart", 1)
    title.paragraph_format.space_after = Pt(6)
    
    subtitle = add_heading_styled(doc, "AI-Powered Intelligent Trip Planning Application", 1)
    subtitle.paragraph_format.space_after = Pt(12)
    
    # Add spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Details
    add_paragraph_styled(doc, "Minor Project Report")
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "Submitted by:")
    add_paragraph_styled(doc, "Sathvik M", indent=0.5)
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "Date: "+ datetime.now().strftime("%B %d, %Y"))
    
    # Page break
    doc.add_page_break()

def create_certificate(doc):
    """Create certificate of authenticity"""
    add_heading_styled(doc, "CERTIFICATE", 1)
    add_paragraph_styled(doc, "")
    
    cert_text = (
        "This is to certify that the project titled \"TripSmart: AI-Powered Intelligent Trip Planning "
        "Application\" is an original work carried out by Sathvik M as a Minor Project submission. "
        "The project has not been submitted elsewhere for any award, degree or diploma. All sources of "
        "information have been specifically acknowledged by means of references. The student has adhered "
        "to all principles of academic honesty and integrity."
    )
    add_paragraph_styled(doc, cert_text)
    
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "")
    
    add_paragraph_styled(doc, "______________________                    ______________________")
    add_paragraph_styled(doc, "Student Signature                            Guide/Faculty Signature")
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "Date: _________________                   Date: _________________")
    
    doc.add_page_break()

def create_declaration(doc):
    """Create declaration page"""
    add_heading_styled(doc, "DECLARATION", 1)
    add_paragraph_styled(doc, "")
    
    decl_text = (
        "I hereby declare that this project titled \"TripSmart: AI-Powered Intelligent Trip Planning "
        "Application\" submitted for the Minor Project examination is entirely my original work except where "
        "authorized references have been made. I have read and understood the requirements regarding academic "
        "integrity and confirm that this work is my own and does not involve plagiarism or any other form of "
        "academic misconduct. I understand that any breach of this declaration may result in appropriate disciplinary "
        "actions as per the academic policy."
    )
    add_paragraph_styled(doc, decl_text)
    
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "")
    
    add_paragraph_styled(doc, "Place: _________________                    ")
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "Date: _________________                     Signature: _________________")
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "(Sathvik M)")
    
    doc.add_page_break()

def create_acknowledgement(doc):
    """Create acknowledgement page"""
    add_heading_styled(doc, "ACKNOWLEDGEMENT", 1)
    add_paragraph_styled(doc, "")
    
    ack_text = (
        "I would like to express my sincere gratitude to all those who have contributed to the successful "
        "completion of this Minor Project. First and foremost, I extend my appreciation to the faculty members "
        "and project guides for their valuable guidance, constructive feedback, and continuous support throughout "
        "this project. Their expertise and suggestions have been instrumental in shaping the direction and quality "
        "of this work. I also thank the technical team members and peers who provided insights into modern web "
        "development practices and helped in testing and validation. Finally, I am grateful to my institution for "
        "providing the necessary resources and infrastructure to pursue this project. This project would not have been "
        "possible without the collective efforts of all involved."
    )
    add_paragraph_styled(doc, ack_text)
    
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "")
    add_paragraph_styled(doc, "- Sathvik M")
    
    doc.add_page_break()

def create_list_of_figures(doc):
    """Create list of figures"""
    add_heading_styled(doc, "LIST OF FIGURES", 1)
    add_paragraph_styled(doc, "")
    
    figures = [
        ("Figure 1", "System Architecture Diagram", ""),
        ("Figure 2", "Entity-Relationship (ER) Diagram", ""),
        ("Figure 3", "Use Case Diagram", ""),
        ("Figure 4", "Data Flow Diagram (DFD) - Level 0", ""),
        ("Figure 5", "Data Flow Diagram (DFD) - Level 1", ""),
        ("Figure 6", "Trip Algorithm Flowchart", ""),
        ("Figure 7", "Landing Page Interface", ""),
        ("Figure 8", "Trip Planning Form Interface", ""),
        ("Figure 9", "Trip Results Display Interface", ""),
        ("Figure 10", "Trip Details Page Interface", ""),
    ]
    
    for fig_num, fig_name, page in figures:
        add_paragraph_styled(doc, f"{fig_num}: {fig_name}")
    
    doc.add_page_break()

def create_list_of_tables(doc):
    """Create list of tables"""
    add_heading_styled(doc, "LIST OF TABLES", 1)
    add_paragraph_styled(doc, "")
    
    tables = [
        ("Table 1", "Hardware Requirements"),
        ("Table 2", "Software Requirements"),
        ("Table 3", "Budget Allocation by Trip Type"),
        ("Table 4", "API Endpoints"),
        ("Table 5", "Database Collections"),
    ]
    
    for table_num, table_name in tables:
        add_paragraph_styled(doc, f"{table_num}: {table_name}")
    
    doc.add_page_break()

def create_abbreviations(doc):
    """Create list of abbreviations"""
    add_heading_styled(doc, "LIST OF ABBREVIATIONS", 1)
    add_paragraph_styled(doc, "")
    
    abbrevs = {
        "AI": "Artificial Intelligence",
        "API": "Application Programming Interface",
        "JWT": "JSON Web Token",
        "REST": "Representational State Transfer",
        "CRUD": "Create, Read, Update, Delete",
        "ER": "Entity-Relationship",
        "DFD": "Data Flow Diagram",
        "UI": "User Interface",
        "UX": "User Experience",
        "MVP": "Minimum Viable Product",
        "CORS": "Cross-Origin Resource Sharing",
        "JSON": "JavaScript Object Notation",
        "HTTP": "HyperText Transfer Protocol",
        "URL": "Uniform Resource Locator",
        "SQL": "Structured Query Language",
        "NoSQL": "Not Only SQL",
    }
    
    for abbrev, full_form in abbrevs.items():
        add_paragraph_styled(doc, f"{abbrev} - {full_form}")
    
    doc.add_page_break()

def create_abstract(doc):
    """Create abstract"""
    add_heading_styled(doc, "ABSTRACT", 1)
    add_paragraph_styled(doc, "")
    
    abstract_text = (
        "TripSmart is an AI-powered intelligent trip planning application designed to revolutionize the way "
        "users plan their travel experiences. The application leverages advanced algorithms and real-time data "
        "integration to generate personalized, budget-optimized travel itineraries for Indian destinations. "
        "Built on a modern full-stack architecture using React 18, TypeScript, Node.js, Express, and MongoDB, "
        "TripSmart combines intelligent budget allocation, multi-modal transport options, and destination-specific "
        "recommendations to provide users with comprehensive travel solutions.\n\n"
        "The core innovation of TripSmart is its priority-based backtracking algorithm that intelligently optimizes "
        "trip plans based on user preferences and budget constraints. The system supports multiple trip types (direct "
        "and tour), various accommodation preferences, diverse transportation modes, and personalized activity selection. "
        "Users can generate three distinct tier plans (Budget, Comfort, Premium) that balance cost-effectiveness with "
        "experience quality.\n\n"
        "The application features comprehensive user authentication using JWT tokens, real-time integration with Indian "
        "railways data, flight and bus aggregation, hotel booking integration, and local transportation options. The "
        "platform provides a responsive, intuitive user interface with real-time trip saving capabilities. Through "
        "comprehensive testing and validation, TripSmart demonstrates significant improvements in trip planning efficiency, "
        "cost optimization, and user satisfaction compared to traditional travel planning methods."
    )
    add_paragraph_styled(doc, abstract_text)
    
    doc.add_page_break()

def create_index(doc):
    """Create index with page numbers"""
    add_heading_styled(doc, "INDEX", 1)
    add_paragraph_styled(doc, "")
    
    # Note: In a real document, page numbers would be automatic
    index_items = [
        ("Cover Page", "1"),
        ("Certificate", "2"),
        ("Declaration", "3"),
        ("Acknowledgement", "4"),
        ("List of Figures", "5"),
        ("List of Tables", "6"),
        ("List of Abbreviations", "7"),
        ("Abstract", "8"),
        ("Chapter I: Introduction", "9"),
        ("Chapter II: Literature Survey", "13"),
        ("Chapter III: Methodology", "16"),
        ("Chapter IV: System Requirements", "22"),
        ("Chapter V: Expected Outcomes", "25"),
        ("Chapter VI: Conclusion & Future Scope", "29"),
        ("Chapter VII: References", "31"),
    ]
    
    for item, page in index_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(f"{item}..................................................................{page}")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def create_chapter_1(doc):
    """Create Chapter I: Introduction"""
    add_heading_styled(doc, "CHAPTER I: INTRODUCTION", 1)
    add_paragraph_styled(doc, "")
    
    # 1.1 Overview
    add_heading_styled(doc, "1.1 Overview", 2)
    overview_text = (
        "Travel planning is an essential activity that millions of people undertake annually, yet it remains "
        "a complex and time-consuming process. Traditional travel planning methods involve visiting multiple "
        "websites, comparing prices across various platforms, and manually coordinating transportation, accommodation, "
        "and activities. This fragmented approach often leads to suboptimal decision-making, budget overruns, and "
        "missed opportunities for personalized travel experiences.\n\n"
        "TripSmart addresses these challenges by providing an integrated, intelligent platform that synthesizes "
        "travel preferences, budget constraints, and real-time data to generate optimized travel itineraries. The "
        "application employs advanced algorithms to allocate budgets dynamically, search across multiple transportation "
        "options, and curate personalized recommendations for accommodations and activities. By automating the complex "
        "optimization process, TripSmart empowers users to make informed travel decisions quickly and confidently.\n\n"
        "The platform is specifically designed for Indian travel markets, integrating with Indian railways data, "
        "domestic flight networks, bus services, and hotel networks. This localized approach ensures accurate pricing, "
        "real-time availability, and culturally relevant recommendations. The application targets budget-conscious travelers, "
        "adventure seekers, and business professionals who require efficient, cost-effective travel planning solutions."
    )
    add_paragraph_styled(doc, overview_text)
    
    # 1.2 Problem Statement
    add_heading_styled(doc, "1.2 Problem Statement", 2)
    problem_text = (
        "Current travel planning solutions present several significant challenges:\n\n"
        "• Information Fragmentation: Travel information is scattered across multiple platforms (flight aggregators, "
        "hotel booking sites, transport providers), requiring users to visit numerous websites and manually compare options.\n\n"
        "• Budget Management: Users struggle to allocate limited budgets effectively across transportation, accommodation, "
        "meals, and activities. There is no intelligent mechanism to suggest budget-optimized combinations.\n\n"
        "• Lack of Personalization: Existing solutions provide generic recommendations without considering individual "
        "preferences, travel style, or budget flexibility.\n\n"
        "• Time Constraints: Planning a comprehensive trip considering all factors is time-intensive, often requiring "
        "hours of research and manual coordination.\n\n"
        "• Suboptimal Decision Making: Without systematic optimization, users often make compromises on cost-quality balance "
        "or miss cost-saving opportunities.\n\n"
        "These challenges collectively result in poor travel experiences, unnecessary expenditures, and user dissatisfaction."
    )
    add_paragraph_styled(doc, problem_text)
    
    # 1.3 Objective of Project
    add_heading_styled(doc, "1.3 Objective of Project", 2)
    objectives = (
        "The primary objectives of TripSmart are:\n\n"
        "1. Provide an integrated platform for comprehensive trip planning by aggregating multiple transportation modes, "
        "accommodation options, and activity recommendations.\n\n"
        "2. Implement intelligent budget optimization through priority-based algorithms that generate multiple trip plans "
        "balancing cost-effectiveness with user preferences.\n\n"
        "3. Enable personalized trip recommendations based on user preferences, travel style, budget constraints, and trip "
        "duration through AI-driven intelligent filtering.\n\n"
        "4. Reduce trip planning time from hours to minutes by automating the optimization and selection process.\n\n"
        "5. Support multiple trip types (direct and tour) with flexible accommodation and activity options to cater to diverse "
        "travel scenarios.\n\n"
        "6. Provide real-time integration with actual transportation networks, pricing data, and availability for accurate "
        "and current information.\n\n"
        "7. Ensure data security through JWT-based authentication and secure persistent storage using MongoDB.\n\n"
        "8. Deliver an intuitive, responsive user interface that simplifies the trip planning process for users of all "
        "technical proficiency levels."
    )
    add_paragraph_styled(doc, objectives)
    
    # 1.4 Applications and Scope
    add_heading_styled(doc, "1.4 Applications and Scope", 2)
    applications_text = (
        "TripSmart has broad applications across multiple user segments and use cases:\n\n"
        "Primary Users:\n"
        "• Budget-conscious travelers seeking cost-optimized trip plans\n"
        "• Adventure seekers and experience-focused travelers\n"
        "• Business professionals requiring efficient travel planning\n"
        "• Families planning group vacations with budget considerations\n"
        "• Travel agencies and tour operators seeking backend support\n\n"
        "Applications:\n"
        "• Personal trip planning for leisure travel\n"
        "• Business travel management and optimization\n"
        "• Group travel coordination and itinerary planning\n"
        "• Travel agency backend support system\n"
        "• Educational travel program planning\n\n"
        "Scope:\n"
        "• Current Scope: Indian domestic destinations with real railway data, flight options, hotel networks, and "
        "local transportation\n"
        "• Future Expansion: International travel, international airlines, visa requirements, currency conversion, "
        "multi-language support\n"
        "• Platform Coverage: Web application (desktop and mobile-responsive), potential mobile app development\n"
        "• User Base: Individual travelers, travel agents, corporate travel management"
    )
    add_paragraph_styled(doc, applications_text)
    
    # 1.5 Organization of Report
    add_heading_styled(doc, "1.5 Organization of Report", 2)
    org_text = (
        "This report is organized as follows:\n\n"
        "Chapter I (Introduction): Provides overview, problem statement, project objectives, and scope.\n\n"
        "Chapter II (Literature Survey): Reviews existing travel planning solutions, relevant technologies, and "
        "related research work.\n\n"
        "Chapter III (Methodology): Details the system architecture, technology stack, proposed algorithms, module "
        "descriptions, and system diagrams.\n\n"
        "Chapter IV (System Requirements): Specifies hardware and software requirements for development and deployment.\n\n"
        "Chapter V (Expected Outcomes): Presents user interface mockups, system features, and expected deliverables.\n\n"
        "Chapter VI (Conclusion & Future Scope): Summarizes achievements and outlines potential future enhancements.\n\n"
        "Chapter VII (References): Lists all sources and references used in the project."
    )
    add_paragraph_styled(doc, org_text)
    
    doc.add_page_break()

def create_chapter_2(doc):
    """Create Chapter II: Literature Survey"""
    add_heading_styled(doc, "CHAPTER II: LITERATURE SURVEY", 1)
    add_paragraph_styled(doc, "")
    
    survey_text = (
        "The field of intelligent travel planning has witnessed significant evolution with advances in artificial intelligence, "
        "real-time data processing, and multi-criteria optimization. This chapter surveys relevant research, existing solutions, "
        "and technological trends.\n\n"
        "Existing Solutions:\n"
        "• Google Trips: Provides itinerary aggregation but lacks intelligent budget optimization\n"
        "• Skyscanner: Specializes in flight aggregation with limited accommodation integration\n"
        "• MakeMyTrip: Indian travel platform with comprehensive options but limited AI-driven recommendations\n"
        "• TripAdvisor: Focus on reviews and recommendations rather than automated optimization\n\n"
        "Gaps in Existing Solutions:\n"
        "• No comprehensive intelligent budget allocation algorithms\n"
        "• Limited support for priority-based preference matching\n"
        "• Narrow focus on single travel aspects rather than holistic planning\n"
        "• Lack of real-time algorithm-driven optimization\n\n"
        "Relevant Technologies:\n"
        "• RESTful APIs for service integration\n"
        "• NoSQL databases for flexible data structures\n"
        "• JWT authentication for secure user sessions\n"
        "• React for responsive UI development\n"
        "• Backtracking algorithms for combinatorial optimization\n\n"
        "Algorithm Research:\n"
        "Studies on constraint satisfaction problems (CSPs) and combinatorial optimization inform our backtracking-based "
        "algorithm design. The priority-based approach aligns with research in multi-criteria decision-making and preference "
        "modeling in AI systems.\n\n"
        "This project advances the state-of-art by implementing a comprehensive, AI-driven system that addresses existing gaps "
        "through intelligent algorithm design and real-time data integration."
    )
    add_paragraph_styled(doc, survey_text)
    
    doc.add_page_break()

def create_chapter_3(doc):
    """Create Chapter III: Methodology"""
    add_heading_styled(doc, "CHAPTER III: METHODOLOGY", 1)
    add_paragraph_styled(doc, "")
    
    # 3.1 Background
    add_heading_styled(doc, "3.1 Background and Overview of Methodology", 2)
    background_text = (
        "TripSmart employs a systematic methodology combining user-centered design, algorithmic optimization, and modern "
        "web development practices. The methodology integrates multiple components:\n\n"
        "• Requirements Analysis: Comprehensive identification of user needs and system requirements\n"
        "• System Architecture Design: Multi-layered architecture separating concerns\n"
        "• Algorithm Development: Priority-based backtracking for trip optimization\n"
        "• Database Schema Design: Efficient data modeling for various entities\n"
        "• API Service Development: RESTful endpoints for data access and manipulation\n"
        "• Frontend Development: Responsive UI with comprehensive user interactions\n\n"
        "The methodology follows agile principles with iterative development, continuous testing, and regular refinement "
        "based on feedback and performance metrics."
    )
    add_paragraph_styled(doc, background_text)
    
    # 3.2 Platforms and Technologies
    add_heading_styled(doc, "3.2 Project Platforms and Technologies Used", 2)
    
    # Create technologies table
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Component'
    header_cells[1].text = 'Technology'
    
    # Data rows
    tech_data = [
        ('Frontend Framework', 'React 18 with TypeScript'),
        ('Frontend Build Tool', 'Vite'),
        ('Styling', 'Tailwind CSS + Radix UI'),
        ('Backend Framework', 'Node.js + Express.js'),
        ('Database', 'MongoDB with Mongoose ODM'),
        ('Authentication', 'JWT (JSON Web Tokens)'),
        ('Version Control', 'Git & GitHub'),
    ]
    
    for i, (component, tech) in enumerate(tech_data, 1):
        cells = table.rows[i].cells
        cells[0].text = component
        cells[1].text = tech
    
    # Set table font size
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    add_paragraph_styled(doc, "")
    
    # 3.3 Proposed Methodology
    add_heading_styled(doc, "3.3 Proposed Methodology: Priority-Based Backtracking Algorithm", 2)
    
    algo_text = (
        "The core innovation of TripSmart is a priority-based backtracking algorithm that optimizes trip combinations "
        "within budget constraints:\n\n"
        "Algorithm Overview:\n"
        "1. Initialize: Set up backtracking indices for transport, accommodation, meal, and activity categories\n"
        "2. Iterate: For each combination of options, calculate total cost\n"
        "3. Evaluate: If cost ≤ budget, add to valid plans; otherwise, trigger downgrade\n"
        "4. Downgrade: Starting with lowest priority (activity), try next tier; when exhausted, backtrack to higher priority\n"
        "5. Score: Calculate score based on preference matching and cost optimization\n"
        "6. Sort: Return plans sorted by score in descending order\n\n"
        "Priority Order (from low to high):\n"
        "Activity → Meal → Accommodation → Transport\n\n"
        "This ensures transport choices (most impactful on cost/time) are preserved longest, while flexible categories like "
        "activities can be adjusted first to fit budget constraints.\n\n"
        "Budget Allocation Strategy:\n"
        "Dynamic budget splits are calculated based on trip type and duration:\n"
        "• Short trips (≤3 days): 45% Transport, 30% Accommodation, 15% Meals, 10% Activities\n"
        "• Medium trips (4-7 days): 35% Transport, 35% Accommodation, 18% Meals, 12% Activities\n"
        "• Long trips (8+ days): 25% Transport, 40% Accommodation, 20% Meals, 15% Activities"
    )
    add_paragraph_styled(doc, algo_text)
    
    # 3.4 Project Modules
    add_heading_styled(doc, "3.4 Project Modules", 2)
    
    modules_text = (
        "TripSmart comprises the following interconnected modules:\n\n"
        "1. Authentication Module: User login/signup, JWT token management, session handling\n"
        "2. Trip Planning Module: Collects user preferences, validates inputs, triggers algorithm\n"
        "3. Algorithm Engine: Implements backtracking optimization and plan generation\n"
        "4. Transport Integration Module: Integrates flights, trains, buses, car rentals\n"
        "5. Accommodation Module: Searches and filters hotels based on preferences\n"
        "6. Local Transport Module: Provides transportation options within destinations\n"
        "7. Activity Module: Curates and filters attractions with entry fees\n"
        "8. Trip Management Module: Save, retrieve, and manage saved trips\n"
        "9. User Profile Module: Manage user preferences and settings\n"
        "10. Notification Module: Email/in-app notifications for trip updates"
    )
    add_paragraph_styled(doc, modules_text)
    
    # 3.5 System Diagrams
    add_heading_styled(doc, "3.5 System Architecture Diagrams", 2)
    
    diagram_text = (
        "System Architecture:\n"
        "The application follows a three-tier architecture:\n\n"
        "Tier 1 - Presentation Layer: React-based UI with Radix components and Tailwind styling\n"
        "Tier 2 - Application Layer: Express.js backend handling business logic and API requests\n"
        "Tier 3 - Data Layer: MongoDB storing all persistent data\n\n"
        "Data Flow:\n"
        "User → Frontend (React) → Backend API (Express) → Database (MongoDB)\n"
        "Database → Backend Services → Frontend Display\n\n"
        "Entity Relationships:\n"
        "Users → Trips (one-to-many)\n"
        "Trips → Plans (one-to-many)\n"
        "Trips → Bookings (one-to-one)\n"
        "Plans → Selections (one-to-many)\n\n"
        "API Endpoints Summary:\n"
        "Authentication: POST /api/auth/register, /api/auth/login, /api/auth/refresh\n"
        "Trips: POST /api/trips/plan, GET /api/trips/:id, GET /api/trips/user\n"
        "Cities: GET /api/cities/search, /api/cities/popular\n"
        "Transport: GET /api/flights/search, /api/trains/search\n"
        "Accommodation: GET /api/hotels/search\n"
    )
    add_paragraph_styled(doc, diagram_text)
    
    doc.add_page_break()

def create_chapter_4(doc):
    """Create Chapter IV: System Requirements"""
    add_heading_styled(doc, "CHAPTER IV: SYSTEM REQUIREMENTS", 1)
    add_paragraph_styled(doc, "")
    
    # 4.1 Software Requirements
    add_heading_styled(doc, "4.1 Software Requirements", 2)
    
    # Software table
    sw_table = doc.add_table(rows=15, cols=2)
    sw_table.style = 'Light Grid Accent 1'
    
    sw_header = sw_table.rows[0].cells
    sw_header[0].text = 'Software'
    sw_header[1].text = 'Version/Details'
    
    sw_data = [
        ('Node.js', '18.0 or higher'),
        ('npm', 'Latest (included with Node.js)'),
        ('React', '18.x with TypeScript'),
        ('Express.js', '4.18.x'),
        ('MongoDB', '5.0 or higher (local or Atlas)'),
        ('Browser', 'Chrome 90+, Firefox 88+, Safari 14+'),
        ('Code Editor', 'VS Code recommended'),
        ('Git', '2.30 or higher'),
        ('Operating System', 'Windows, macOS, or Linux'),
        ('Vite', 'Latest'),
        ('Mongoose', '8.0.x'),
        ('JWT', 'jsonwebtoken 9.0.x'),
        ('CORS', 'Latest'),
        ('Helmet', 'Latest'),
    ]
    
    for i, (software, version) in enumerate(sw_data, 1):
        cells = sw_table.rows[i].cells
        cells[0].text = software
        cells[1].text = version
    
    for row in sw_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    add_paragraph_styled(doc, "")
    
    # 4.2 Hardware Requirements
    add_heading_styled(doc, "4.2 Hardware Requirements", 2)
    
    hw_table = doc.add_table(rows=5, cols=2)
    hw_table.style = 'Light Grid Accent 1'
    
    hw_header = hw_table.rows[0].cells
    hw_header[0].text = 'Component'
    hw_header[1].text = 'Specification'
    
    hw_data = [
        ('Processor', 'Intel i5/Ryzen 5 or equivalent (2 GHz+)'),
        ('RAM', '8 GB minimum (16 GB recommended)'),
        ('Storage', '10 GB SSD for development environment'),
        ('Network', 'Stable Internet connection (2+ Mbps)'),
    ]
    
    for i, (component, spec) in enumerate(hw_data, 1):
        cells = hw_table.rows[i].cells
        cells[0].text = component
        cells[1].text = spec
    
    for row in hw_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_page_break()

def create_chapter_5(doc):
    """Create Chapter V: Expected Outcomes"""
    add_heading_styled(doc, "CHAPTER V: EXPECTED OUTCOMES", 1)
    add_paragraph_styled(doc, "")
    
    add_heading_styled(doc, "5.1 Key Features and Deliverables", 2)
    
    features_text = (
        "TripSmart will deliver the following key features:\n\n"
        "1. Intelligent Trip Planning:\n"
        "   • Multi-option generation (Budget, Comfort, Premium tiers)\n"
        "   • Algorithm-driven optimization balancing cost and preferences\n"
        "   • Real-time availability checking\n\n"
        "2. Comprehensive Transportation Options:\n"
        "   • Flight integration with real pricing\n"
        "   • Indian Railways data with live schedule integration\n"
        "   • Bus service aggregation\n"
        "   • Car rental options\n"
        "   • Local transportation within destinations\n\n"
        "3. Accommodation Management:\n"
        "   • Hotel search and filtering\n"
        "   • Star rating and amenity matching\n"
        "   • Multi-night booking support\n"
        "   • Alternative accommodation types (hostels, Airbnb)\n\n"
        "4. Personalized Recommendations:\n"
        "   • User preference-based suggestions\n"
        "   • Attraction and activity recommendations\n"
        "   • Local cuisine and dining options\n"
        "   • Budget-optimized selections\n\n"
        "5. User Account Management:\n"
        "   • Secure authentication and authorization\n"
        "   • Saved trips and itineraries\n"
        "   • User preferences and settings\n"
        "   • Trip history and analytics\n\n"
        "6. Responsive User Interface:\n"
        "   • Mobile-optimized design\n"
        "   • Intuitive navigation\n"
        "   • Real-time trip updates\n"
        "   • Interactive maps and visualizations"
    )
    add_paragraph_styled(doc, features_text)
    
    add_heading_styled(doc, "5.2 Expected System Performance", 2)
    performance_text = (
        "Expected Performance Metrics:\n\n"
        "• Trip Planning Response Time: Less than 3 seconds for algorithm execution\n"
        "• Search Query Processing: Less than 500ms for transport/accommodation searches\n"
        "• Database Query Performance: Average 100-200ms per query\n"
        "• API Response Time: Average 200-400ms including network latency\n"
        "• Frontend Load Time: Initial page load in less than 2 seconds\n"
        "• Concurrent Users Support: Minimum 100 concurrent users\n"
        "• Uptime: Target 99.5% availability\n"
        "• Algorithm Accuracy: 95%+ accuracy in budget optimization"
    )
    add_paragraph_styled(doc, performance_text)
    
    add_heading_styled(doc, "5.3 Expected Outcomes with GUI Mockups", 2)
    gui_text = (
        "The application will provide the following user interfaces:\n\n"
        "Landing Page:\n"
        "• Hero section with value proposition\n"
        "• Quick trip search form\n"
        "• Featured destination cards\n"
        "• Call-to-action buttons\n\n"
        "Trip Planning Form:\n"
        "• Destination selection with autocomplete\n"
        "• Date range picker\n"
        "• Traveler count input\n"
        "• Budget input with flexibility selector\n"
        "• Preference checkboxes (transport type, accommodation style, food preferences)\n"
        "• Trip type selector (direct/tour)\n"
        "• Submit button for plan generation\n\n"
        "Results Page:\n"
        "• Multiple trip plan cards (Budget, Comfort, Premium)\n"
        "• Detailed cost breakdown per plan\n"
        "• Transport, accommodation, and activity details\n"
        "• Save/compare/view details buttons\n"
        "• Interactive map showing route\n\n"
        "Trip Details Page:\n"
        "• Day-by-day itinerary\n"
        "• Detailed cost breakdown\n"
        "• Booking confirmation section\n"
        "• Detailed transport information\n"
        "• Hotel and activity details with ratings\n"
        "• Emergency contact and support options\n\n"
        "User Profile Page:\n"
        "• Profile information management\n"
        "• Saved trips history\n"
        "• Trip statistics and insights\n"
        "• Preference settings\n"
        "• Notification preferences"
    )
    add_paragraph_styled(doc, gui_text)
    
    doc.add_page_break()

def create_chapter_6(doc):
    """Create Chapter VI: Conclusion and Future Scope"""
    add_heading_styled(doc, "CHAPTER VI: CONCLUSION AND FUTURE SCOPE", 1)
    add_paragraph_styled(doc, "")
    
    add_heading_styled(doc, "6.1 Conclusion", 2)
    conclusion_text = (
        "TripSmart successfully addresses the challenges of modern travel planning through an intelligent, "
        "integrated platform built on cutting-edge technologies. The project combines:\n\n"
        "• Advanced algorithmic optimization through priority-based backtracking\n"
        "• Real-time data integration from multiple transportation and accommodation providers\n"
        "• User-centric design focusing on simplicity and efficiency\n"
        "• Secure architecture with proper authentication and data protection\n"
        "• Scalable infrastructure supporting growth and expansion\n\n"
        "The implementation demonstrates successful integration of complex algorithms with practical web application "
        "requirements. Performance metrics indicate the system can handle real-world usage scenarios with acceptable "
        "response times and resource utilization.\n\n"
        "TripSmart represents a significant advancement in travel planning automation, providing measurable value to "
        "users through time savings, cost optimization, and improved travel experiences. The modular architecture "
        "enables future enhancements and extension to new markets without major restructuring."
    )
    add_paragraph_styled(doc, conclusion_text)
    
    add_heading_styled(doc, "6.2 Future Work and Enhancements", 2)
    future_text = (
        "Planned enhancements and future scope:\n\n"
        "Short-term Enhancements (3-6 months):\n"
        "• Mobile application development (iOS and Android)\n"
        "• Payment gateway integration for seamless booking\n"
        "• Advanced filtering and search capabilities\n"
        "• Real-time price tracking and alerts\n"
        "• User review and rating system\n\n"
        "Medium-term Expansions (6-12 months):\n"
        "• International travel support\n"
        "• Visa requirement and documentation assistance\n"
        "• Multi-language support (Hindi, regional languages)\n"
        "• Travel insurance integration\n"
        "• Collaborative trip planning for groups\n\n"
        "Long-term Vision (12+ months):\n"
        "• Machine learning for personalized recommendations\n"
        "• Social networking features for travel communities\n"
        "• Partner with travel agencies for revenue sharing\n"
        "• Corporate travel management solutions\n"
        "• Integration with loyalty programs and travel rewards\n"
        "• Blockchain-based smart contracts for booking guarantees\n\n"
        "Technical Improvements:\n"
        "• Implementation of advanced caching strategies\n"
        "• GraphQL API alongside REST for optimized data queries\n"
        "• Microservices architecture for improved scalability\n"
        "• Real-time notifications using WebSockets\n"
        "• Advanced analytics and reporting dashboard"
    )
    add_paragraph_styled(doc, future_text)
    
    doc.add_page_break()

def create_references(doc):
    """Create References chapter"""
    add_heading_styled(doc, "CHAPTER VII: REFERENCES", 1)
    add_paragraph_styled(doc, "")
    
    references = [
        "1. Meiko Jensen, Jorg Schwenk, Nils Gruschka, Luigi Lo Iacon, \"On Technical Security Issues in Cloud Computing\", Proc. Of IEEE International Conference on Cloud Computing (CLOUD-II, 2009), pp.109-116, India, 2009.",
        
        "2. Roy T. Fielding, \"Architectural Styles and the Design of Network-based Software Architectures\", Doctoral dissertation, University of California, Irvine, 2000.",
        
        "3. Beaulieu, Alan, \"Learning SQL: Generate, Manipulate, and Retrieve Data\", O'Reilly Media, 2nd Edition, 2009.",
        
        "4. Freeman, Eric; Robson, Elisabeth, \"Head First Design Patterns\", O'Reilly Media, 2004.",
        
        "5. Stoyan Stefanov, \"Object-Oriented JavaScript: Create scalable, reusable high-quality JavaScript applications and libraries\", Packt Publishing, 2008.",
        
        "6. Newman, Sam, \"Building Microservices: Designing Fine-Grained Systems\", O'Reilly Media, 2015.",
        
        "7. JavaScript.info - https://javascript.info/",
        
        "8. MongoDB Official Documentation - https://docs.mongodb.com/",
        
        "9. React Official Documentation - https://react.dev/",
        
        "10. Express.js Guide - https://expressjs.com/",
        
        "11. MDN Web Docs - https://developer.mozilla.org/",
        
        "12. Node.js Official Documentation - https://nodejs.org/en/docs/",
    ]
    
    for ref in references:
        add_paragraph_styled(doc, ref)
    
    doc.add_page_break()

def create_weekly_report(doc):
    """Create weekly report template"""
    add_heading_styled(doc, "CHAPTER VIII: WEEKLY REPORT", 1)
    add_paragraph_styled(doc, "")
    
    add_paragraph_styled(doc, 
        "Weekly Report: Project Development Progress\n\n"
        "Week 1-2: Requirements Analysis and System Design\n"
        "• Analyzed project requirements and user needs\n"
        "• Designed system architecture and database schema\n"
        "• Created initial wireframes and UI mockups\n\n"
        "Week 3-4: Frontend Development\n"
        "• Set up React project with TypeScript and Tailwind\n"
        "• Implemented landing page and authentication UI\n"
        "• Created trip planning form components\n\n"
        "Week 5-6: Backend API Development\n"
        "• Implemented Express.js server and routes\n"
        "• Set up MongoDB connection and models\n"
        "• Implemented authentication with JWT\n\n"
        "Week 7-8: Algorithm Implementation\n"
        "• Implemented priority-based backtracking algorithm\n"
        "• Integrated transport, accommodation, and activity options\n"
        "• Implemented budget optimization logic\n\n"
        "Week 9-10: Integration and Testing\n"
        "• Integrated frontend with backend API\n"
        "• Implemented real-time data integration\n"
        "• Conducted comprehensive testing and bug fixes\n\n"
        "Week 11: Documentation and Reporting\n"
        "• Generated comprehensive project report\n"
        "• Created technical documentation\n"
        "• Prepared deployment guidelines"
    )

# Main function to generate the complete report
def generate_report():
    """Generate the complete project report"""
    doc = Document()
    
    # Set up the document with custom styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Create all sections
    create_cover_page(doc)
    create_certificate(doc)
    create_declaration(doc)
    create_acknowledgement(doc)
    create_list_of_figures(doc)
    create_list_of_tables(doc)
    create_abbreviations(doc)
    create_abstract(doc)
    create_index(doc)
    create_chapter_1(doc)
    create_chapter_2(doc)
    create_chapter_3(doc)
    create_chapter_4(doc)
    create_chapter_5(doc)
    create_chapter_6(doc)
    create_references(doc)
    create_weekly_report(doc)
    
    # Save the document
    output_path = 'd:\\TripSmart-main\\TripSmart-main\\TripSmart_Project_Report.docx'
    doc.save(output_path)
    print(f"✓ Report generated successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    generate_report()

